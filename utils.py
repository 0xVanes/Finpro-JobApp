from __future__ import annotations

import io
import logging
import os

import requests
import streamlit as st
from dotenv import find_dotenv, load_dotenv

# ─── KONSTANTA KONFIGURASI (NFR-4.03) ───────────────────────────
# Semua nilai sensitif dari environment variable (.env)
# Tidak ada credential yang di-hardcode di sini
load_dotenv(find_dotenv())

N8N_WEBHOOK_URL: str | None = os.environ.get("N8N_WEBHOOK_URL")

# Basic Auth untuk webhook n8n (opsional — kosongkan di .env kalau
# webhook tidak pakai Basic Auth). Kalau salah satu USER/PASS terisi,
# call_n8n() otomatis menyertakan header Authorization di tiap request.
N8N_BASIC_AUTH_USER: str | None = os.environ.get("N8N_BASIC_AUTH_USER")
N8N_BASIC_AUTH_PASS: str | None = os.environ.get("N8N_BASIC_AUTH_PASS")

# Satu logger per modul — tidak pakai print() untuk error internal
logger = logging.getLogger(__name__)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)

if not N8N_WEBHOOK_URL:
    # Jangan biarkan aplikasi diam-diam gagal di call_n8n() tanpa
    # penjelasan — beri tahu sejak awal lewat log kalau .env belum diisi.
    logger.warning(
        "N8N_WEBHOOK_URL tidak ditemukan di environment. "
        "Pastikan file .env berisi N8N_WEBHOOK_URL=<url webhook n8n kamu>."
    )

# Timeout (detik) per jenis request sesuai NFR-2
#
# Catatan TIMEOUT_SQL: workflow n8n untuk mode "search" saat ini masih
# memakai AI Agent (LangChain Agent + tool SQL), bukan direct query.
# Nilai 5s sesuai target NFR-2.02 TIDAK realistis untuk pipeline
# berbasis LLM (butuh waktu untuk agent menyusun & menjalankan query).
# Dinaikkan ke 20s supaya tidak timeout prematur di kondisi normal.
# (Kalau nanti mode search disederhanakan jadi direct SQL tanpa AI
# Agent, nilai ini bisa diturunkan lagi mendekati target asli.)
TIMEOUT_CHAT:  int = 30   # NFR-2.01 RAG chat
TIMEOUT_SQL:   int = 20   # Lihat catatan di atas
TIMEOUT_CV:    int = 150  # mode cv_match sekarang = 5 AI Agent berantai
                          # (CV Matcher -> Career Path -> Gap Analysis ->
                          # Certification -> Roadmap). Tiap agent bisa
                          # makan 10-20 detik (LLM + tool call), jadi
                          # total realistis 50-100+ detik. 60s (nilai
                          # sebelumnya) terbukti kurang di percobaan
                          # nyata — dinaikkan ke 150s dengan margin aman.
TIMEOUT_OTHER: int = 20   # Default untuk mode lainnya

# Format file CV yang didukung (NFR-3.02)
SUPPORTED_CV_TYPES: tuple[str, ...] = ("pdf", "docx")


# ─── HELPER: SESSION STATE ───────────────────────────────────────

_STATE_DEFAULTS: dict[str, object] = {
    "chat_history": [],
}


def init_session_state() -> None:
    """
    Inisialisasi session_state dengan nilai default (idempoten).

    Aman dipanggil berkali-kali — hanya set nilai default kalau
    key belum ada di session_state (tidak menimpa data yang sudah
    ada, misal chat_history yang sedang berjalan).
    """
    for key, default in _STATE_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = list(default) if isinstance(default, list) else default


# ─── HELPER: CSS ─────────────────────────────────────────────────

def load_css(path: str = "styles.css") -> None:
    """
    Muat dan inject CSS eksternal ke halaman Streamlit aktif.

    PENTING: selalu pakai encoding="utf-8" eksplisit — tanpa ini,
    Windows akan pakai encoding default (cp1252) dan merusak karakter
    non-ASCII di CSS (emoji, simbol panah, dll jadi mojibake).

    Args:
        path: Path ke file CSS, relatif terhadap direktori kerja
              Streamlit (default: "styles.css" di root project).
    """
    try:
        with open(path, "r", encoding="utf-8") as f:
            css = f.read()
        st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        logger.warning("styles.css tidak ditemukan di path: %s", path)


# ─── HELPER: N8N ────────────────────────────────────────────────

def call_n8n(payload: dict, timeout: int = TIMEOUT_OTHER) -> dict:
    """
    Kirim request ke N8N webhook dan kembalikan respons JSON.

    Semua exception ditangkap dan dikembalikan sebagai
    ``{"error": pesan}`` sehingga caller tinggal cek key "error".
    Tidak ada traceback Python yang bocor ke UI (NFR-3.02).

    Args:
        payload: Dict yang akan di-serialize menjadi JSON body.
        timeout: Batas waktu tunggu dalam detik.

    Returns:
        Dict hasil JSON dari N8N, atau ``{"error": str}`` jika gagal.
    """
    if not N8N_WEBHOOK_URL:
        return {"error": (
            "⚙️ N8N_WEBHOOK_URL belum dikonfigurasi. "
            "Hubungi administrator aplikasi."
        )}

    # Basic Auth opsional — hanya disertakan kalau kredensial terisi di .env
    auth = None
    if N8N_BASIC_AUTH_USER and N8N_BASIC_AUTH_PASS:
        auth = (N8N_BASIC_AUTH_USER, N8N_BASIC_AUTH_PASS)

    try:
        response = requests.post(
            N8N_WEBHOOK_URL,
            json=payload,
            timeout=timeout,
            auth=auth,
        )
        response.raise_for_status()
        return response.json()                          # type: ignore[no-any-return]

    except requests.exceptions.Timeout:
        logger.warning("N8N timeout setelah %ds — payload mode: %s",
                       timeout, payload.get("mode", "?"))
        return {"error": (
            f"⏱️ Waktu habis ({timeout}s). "
            "Server sedang sibuk, silakan coba lagi."
        )}

    except requests.exceptions.ConnectionError:
        logger.error("Tidak dapat menjangkau N8N: %s", N8N_WEBHOOK_URL)
        return {"error": (
            "🔌 Tidak dapat terhubung ke layanan AI. "
            "Periksa koneksi internet kamu."
        )}

    except requests.exceptions.HTTPError as exc:
        status = exc.response.status_code if exc.response is not None else "?"
        logger.error("N8N HTTP error %s — payload: %s", status, payload)
        if status == 401:
            return {"error": (
                "🔒 Autentikasi ke server ditolak (401). "
                "Cek N8N_BASIC_AUTH_USER/N8N_BASIC_AUTH_PASS di .env."
            )}
        return {"error": f"🚫 Server merespons dengan kode {status}. Coba lagi."}

    except ValueError as exc:
        # JSON decode error
        logger.error("N8N respons bukan JSON valid: %s", exc)
        return {"error": "⚠️ Respons dari server tidak dapat dibaca. Hubungi administrator."}

    except Exception as exc:                            # noqa: BLE001
        logger.exception("N8N call gagal tidak terduga: %s", exc)
        return {"error": "❌ Terjadi kesalahan tak terduga. Hubungi administrator."}


# ─── HELPER: CV — EKSTRAKSI TEKS + FALLBACK OCR ─────────────────
#
# OCR dipakai sebagai FALLBACK saja — hanya dijalankan kalau ekstraksi
# teks normal tidak menghasilkan teks yang cukup (mis. CV hasil scan
# atau desain visual yang teksnya berbentuk gambar). Backend OCR: OpenAI
# Vision (gpt-4o-mini) — tidak perlu instalasi Tesseract di sistem operasi,
# cukup pakai OPENAI_API_KEY yang sudah dipakai fitur lain di project ini.

MIN_CHARS_BEFORE_OCR: int = 20
PDF_RENDER_DPI: int = 200
OCR_MAX_TOKENS: int = 2000     # cukup untuk ~1 halaman CV padat teks
OCR_TIMEOUT_SECONDS: int = 30  # jangan sampai 1 gambar bikin request menggantung lama


def extract_cv_text(uploaded_file) -> str:
    """
    Ekstrak teks mentah dari file CV yang diupload.

    Mendukung PDF dan DOCX, termasuk fallback OCR untuk CV yang
    berupa hasil scan/gambar. File hanya dibaca ke memory (NFR-5.01),
    tidak disimpan ke disk.

    Args:
        uploaded_file: Objek UploadedFile dari st.file_uploader().

    Returns:
        String teks hasil ekstraksi, atau string kosong jika gagal.
    """
    try:
        raw_bytes = uploaded_file.read()
        file_ext = uploaded_file.name.rsplit(".", 1)[-1].lower()

        if file_ext == "pdf":
            return _extract_pdf(raw_bytes)
        elif file_ext == "docx":
            return _extract_docx(raw_bytes)
        else:
            logger.warning("Format CV tidak didukung: .%s", file_ext)
            return ""

    except Exception as exc:                            # noqa: BLE001
        logger.error("extract_cv_text gagal: %s", exc)
        return ""


def _ocr_image_bytes(image_bytes: bytes) -> str:
    """
    Jalankan OCR pada satu gambar lewat OpenAI Vision (model gpt-4o-mini).

    Dipilih dibanding Tesseract karena tidak perlu instalasi program
    terpisah di sistem operasi (Tesseract OCR engine) — cukup pakai
    OPENAI_API_KEY yang sudah dipakai fitur lain di project ini.
    Trade-off: butuh koneksi internet & kena biaya API per panggilan
    (kecil, tapi bukan gratis seperti Tesseract lokal).

    Return string kosong kalau OCR tidak tersedia/gagal — tidak pernah
    raise, supaya ekstraksi CV tetap lanjut jalan tanpa OCR kalau ini gagal.
    """
    try:
        import base64
        from openai import OpenAI

        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.warning(
                "OCR dilewati: OPENAI_API_KEY tidak ditemukan di environment."
            )
            return ""

        client = OpenAI(api_key=api_key)
        base64_image = base64.b64encode(image_bytes).decode("utf-8")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Ekstrak SEMUA teks yang terlihat di gambar ini "
                            "apa adanya (verbatim), tanpa komentar atau "
                            "penjelasan tambahan. Kalau gambar tidak "
                            "mengandung teks sama sekali, jawab dengan "
                            "string kosong."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{base64_image}"},
                    },
                ],
            }],
            max_tokens=OCR_MAX_TOKENS,
            timeout=OCR_TIMEOUT_SECONDS,
        )
        return (response.choices[0].message.content or "").strip()

    except ImportError:
        logger.warning(
            "OCR dilewati: package 'openai' tidak terinstall. "
            "Jalankan: pip install openai"
        )
        return ""
    except Exception as exc:
        logger.warning("OCR (OpenAI Vision) gagal diproses: %s", exc)
        return ""


def _extract_pdf(raw: bytes) -> str:
    """Ekstrak teks dari PDF, dengan fallback OCR per halaman jika teks native terlalu sedikit."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf tidak terinstall. Jalankan: pip install pypdf")
        return ""

    try:
        reader = PdfReader(io.BytesIO(raw))
        page_texts: list[str] = []

        for page_number, page in enumerate(reader.pages):
            native_text = (page.extract_text() or "").strip()
            if len(native_text) >= MIN_CHARS_BEFORE_OCR:
                page_texts.append(native_text)
                continue
            ocr_text = _ocr_pdf_page(raw, page_number)
            page_texts.append(ocr_text or native_text)

        return "\n".join(page_texts).strip()

    except Exception as exc:
        logger.error("Gagal membaca PDF: %s", exc)
        return ""


def _ocr_pdf_page(raw_pdf: bytes, page_number: int) -> str:
    """Render satu halaman PDF jadi gambar (via PyMuPDF), lalu OCR."""
    try:
        import fitz  # PyMuPDF

        pdf_doc = fitz.open(stream=raw_pdf, filetype="pdf")
        page = pdf_doc.load_page(page_number)
        zoom = PDF_RENDER_DPI / 72
        pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        return _ocr_image_bytes(pixmap.tobytes("png"))
    except ImportError:
        logger.warning("OCR PDF dilewati: pymupdf tidak terinstall. Jalankan: pip install pymupdf")
        return ""
    except Exception as exc:
        logger.warning("Gagal render halaman %d untuk OCR: %s", page_number, exc)
        return ""


def _extract_docx(raw: bytes) -> str:
    """
    Ekstrak teks dari DOCX: paragraf + tabel + text box, dengan
    fallback OCR gambar tersemat kalau hasil gabungan masih sedikit.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx tidak terinstall. Jalankan: pip install python-docx")
        return ""

    try:
        doc = Document(io.BytesIO(raw))
        text_parts = _collect_docx_text_parts(doc)
        combined_text = "\n".join(text_parts).strip()

        if len(combined_text) >= MIN_CHARS_BEFORE_OCR:
            return combined_text

        ocr_text = _ocr_docx_embedded_images(raw)
        return (combined_text + "\n" + ocr_text).strip()

    except Exception as exc:
        logger.error("Gagal membaca DOCX: %s", exc)
        return ""


def _collect_docx_text_parts(doc) -> list[str]:
    """Kumpulkan teks dari paragraf, tabel, dan text box sebuah dokumen DOCX."""
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs if p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    try:
        import re
        textbox_matches = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", doc.element.xml)
        textbox_text = " ".join(t for t in textbox_matches if t.strip())
        if textbox_text.strip():
            parts.append(textbox_text)
    except Exception as exc:
        logger.warning("Gagal parsing text box DOCX: %s", exc)

    return parts


def _ocr_docx_embedded_images(raw: bytes) -> str:
    """Ekstrak dan OCR semua gambar tersemat (word/media/*) di dalam DOCX."""
    import zipfile

    supported_extensions = (".png", ".jpg", ".jpeg", ".bmp", ".tiff")
    ocr_results: list[str] = []

    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as docx_zip:
            media_files = [
                name for name in docx_zip.namelist()
                if name.startswith("word/media/") and name.lower().endswith(supported_extensions)
            ]
            for media_name in media_files:
                ocr_text = _ocr_image_bytes(docx_zip.read(media_name))
                if ocr_text:
                    ocr_results.append(ocr_text)
    except Exception as exc:
        logger.warning("Gagal ekstrak gambar tersemat dari DOCX: %s", exc)

    return "\n".join(ocr_results)


def validate_cv_upload(uploaded_file) -> tuple[bool, str]:
    """
    Validasi file CV sebelum diproses.

    Args:
        uploaded_file: Objek UploadedFile dari st.file_uploader().

    Returns:
        Tuple (valid: bool, pesan_error: str).
        Jika valid, pesan_error adalah string kosong.
    """
    if uploaded_file is None:
        return False, "Belum ada file yang diupload."

    ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
    if ext not in SUPPORTED_CV_TYPES:
        return False, (
            f"Format '{ext}' tidak didukung. "
            f"Harap upload file PDF atau DOCX."
        )

    max_bytes = 10 * 1024 * 1024
    if uploaded_file.size > max_bytes:
        return False, (
            f"File terlalu besar ({uploaded_file.size // 1024 // 1024} MB). "
            "Maksimal 10 MB."
        )

    return True, ""


# ─── HELPER: RENDER KOMPONEN UI ─────────────────────────────────

def section_lbl(text: str, icon: str = "") -> None:
    """Render label bagian kecil berwarna sebelum kelompok konten."""
    prefix = f"{icon} " if icon else ""
    st.markdown(f'<div class="section-lbl">{prefix}{text}</div>', unsafe_allow_html=True)


def render_skill_chips(matching: list[str], missing: list[str]) -> None:
    """Render chip skill: hijau (cocok, class .chip.has) vs oranye (kurang, class .chip.miss)."""
    chips = (
        "".join(f'<span class="chip has">✓ {s}</span>' for s in matching)
        + "".join(f'<span class="chip miss">✗ {s}</span>' for s in missing)
    )
    if chips:
        st.markdown(chips, unsafe_allow_html=True)


def show_error(message: str) -> None:
    """Tampilkan kotak error ramah pengguna tanpa traceback Python (NFR-3.02)."""
    st.markdown(f'<div class="err-box">{message}</div>', unsafe_allow_html=True)


def job_card(job: dict) -> None:
    """
    Render kartu lowongan (class .job-card di styles.css — cocokkan
    kalau styles.css berubah nama class-nya lagi).

    Menggunakan .get() dengan default di semua field sehingga tidak
    crash jika field tidak ada di respons N8N.

    Args:
        job: Dict berisi data lowongan dari N8N.
    """
    salary = job.get("salary") or job.get("salary_raw") or "Tidak disebutkan"
    if salary in ("None", "null", ""):
        salary = "Tidak disebutkan"

    reason = job.get("relevance_reason") or job.get("reason") or ""
    rank   = job.get("rank", "")

    rank_html = (
        f'<div style="font-size:.7rem;color:#6D5DF6;font-weight:700;'
        f'margin-bottom:2px">{rank}</div>'
    ) if rank else ""

    reason_html = f'<div class="job-reason">💡 {reason}</div>' if reason else ""

    st.markdown(
        f'<div class="job-card">'
        f'{rank_html}'
        f'<h3>💼 {job.get("job_title", "N/A")}</h3>'
        f'<div class="job-company">{job.get("company_name", "N/A")}</div>'
        f'<div class="job-meta">'
        f'📍 {job.get("location", "N/A")} · {job.get("work_type", "N/A")}</div>'
        f'<div style="margin-top:8px">'
        f'<span class="salary-tag">💰 {salary}</span></div>'
        f'{reason_html}'
        f'</div>',
        unsafe_allow_html=True,
    )


# ─── HELPER: INTERVIEW — EKSPOR PDF (FR-6.05) ───────────────────

def _latin1_safe(text: object) -> str:
    """
    Normalisasi teks agar aman dirender core font FPDF (latin-1).

    Karakter tipografis umum (en/em dash, tanda kutip lengkung, panah,
    bullet) diganti padanan ASCII; sisanya di luar latin-1 diganti '?'.
    Mencegah UnicodeEncodeError saat menulis teks berbahasa Indonesia
    atau jawaban pengguna yang memuat simbol/emoji.
    """
    if text is None:
        return ""
    s = str(text)
    replacements = {
        "–": "-", "—": "-", "•": "-", "→": "->", "←": "<-",
        "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...",
        "✓": "v", "✗": "x", "≥": ">=", "≤": "<=",
    }
    for a, b in replacements.items():
        s = s.replace(a, b)
    return s.encode("latin-1", "replace").decode("latin-1")


def build_interview_pdf(
    job: dict,
    cv_profile: dict,
    history: list,
    assessment: dict,
) -> bytes | None:
    """
    Bangun PDF transkrip wawancara + penilaian (FR-6.05).

    Import fpdf di dalam fungsi (lazy) mengikuti pola helper lain di modul
    ini, sehingga halaman lain tidak wajib memiliki fpdf2 terpasang.

    Args:
        job: Lowongan terpilih (job_id, job_title, company_name).
        cv_profile: Profil ringkas kandidat dari response start.
        history: Daftar {question, answer, category}.
        assessment: Hasil penilaian akhir (verdict, score, reasons, ...).

    Returns:
        Bytes PDF, atau None jika fpdf2 belum terinstall / gagal
        (caller menyediakan fallback unduh teks).
    """
    try:
        from fpdf import FPDF
    except ImportError:
        logger.warning(
            "build_interview_pdf dilewati: fpdf2 tidak terinstall. "
            "Jalankan: pip install fpdf2"
        )
        return None

    try:
        pdf = FPDF(format="A4")
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        def write(text: object, height: float = 6) -> None:
            """Tulis satu blok teks lebar penuh, kursor kembali ke margin kiri.

            new_x/new_y eksplisit mencegah error 'Not enough horizontal space'
            akibat kursor tertinggal di margin kanan setelah multi_cell(w=0).
            """
            pdf.multi_cell(0, height, _latin1_safe(text), new_x="LMARGIN", new_y="NEXT")

        # Judul
        pdf.set_font("Helvetica", "B", 16)
        write("Transkrip Simulasi Wawancara", 9)
        pdf.set_font("Helvetica", "", 10)
        write(
            f"Posisi: {job.get('job_title', '-')}  |  "
            f"Perusahaan: {job.get('company_name', '-')}"
        )
        skills = ", ".join(cv_profile.get("key_skills", []) or []) or "-"
        write(f"Skill kandidat: {skills}")
        pdf.ln(3)

        # Penilaian akhir (FR-6.04)
        pdf.set_font("Helvetica", "B", 13)
        write("Hasil Penilaian", 8)
        pdf.set_font("Helvetica", "B", 11)
        verdict = assessment.get("verdict", "-")
        score = assessment.get("score", "-")
        write(f"Verdict: {verdict}   (Skor: {score}/100)", 7)

        # Cakupan wawancara & tingkat keyakinan
        coverage = assessment.get("coverage", {}) or {}
        if coverage or assessment.get("confidence"):
            pdf.set_font("Helvetica", "", 10)
            write(
                f"Cakupan: {coverage.get('explored_count', 0)}/"
                f"{coverage.get('total_points', 0)} poin tergali "
                f"({coverage.get('ratio_pct', 0)}%)  |  "
                f"Keyakinan: {assessment.get('confidence', '-')}"
            )
            if coverage.get("note"):
                write(coverage["note"])
            untested = coverage.get("untested_areas", []) or []
            if untested:
                write("Tidak pernah diuji: " + ", ".join(str(u) for u in untested))
            pdf.ln(1)

        # Rubrik per dimensi — dasar perhitungan skor
        dimensions = assessment.get("dimensions", []) or []
        if dimensions:
            pdf.set_font("Helvetica", "B", 10)
            write("Rubrik (dasar skor):")
            for dim in dimensions:
                pdf.set_font("Helvetica", "", 10)
                write(f"  - {dim.get('name', '-')}: {dim.get('score', '-')}/10")
                if dim.get("evidence"):
                    pdf.set_font("Helvetica", "I", 9)
                    write(f"      bukti: {dim['evidence']}", 5)

        def _write_items(label: str, key: str) -> None:
            """Tulis butir penilaian; mendukung item string maupun {point, evidence}."""
            items = assessment.get(key, []) or []
            if not items:
                return
            pdf.set_font("Helvetica", "B", 10)
            write(f"{label}:")
            for it in items:
                if isinstance(it, dict):
                    point, evidence = it.get("point", ""), it.get("evidence", "")
                else:
                    point, evidence = it, ""
                pdf.set_font("Helvetica", "", 10)
                write(f"  - {point}")
                if evidence:
                    pdf.set_font("Helvetica", "I", 9)
                    write(f"      bukti: {evidence}", 5)

        _write_items("Alasan", "reasons")
        _write_items("Kelebihan", "strengths")
        _write_items("Perlu ditingkatkan", "improvements")
        _write_items("Kesadaran diri (dinilai positif)", "self_awareness")
        pdf.ln(2)

        # Transkrip Q&A (FR-6.05)
        pdf.set_font("Helvetica", "B", 13)
        write("Transkrip Tanya-Jawab", 8)
        per_q = {i: pq for i, pq in enumerate(assessment.get("per_question", []) or [])}
        for i, turn in enumerate(history):
            cat = turn.get("category", "")
            cat_txt = f" [{cat}]" if cat else ""
            pdf.set_font("Helvetica", "B", 10)
            write(f"Q{i + 1}{cat_txt}: {turn.get('question', '')}")
            pdf.set_font("Helvetica", "", 10)
            write(f"Jawaban: {turn.get('answer', '')}")
            fb = per_q.get(i, {})
            if fb.get("feedback"):
                pdf.set_font("Helvetica", "I", 9)
                fb_score = fb.get("score")
                score_txt = f" (nilai {fb_score}/10)" if fb_score is not None else ""
                write(f"Catatan{score_txt}: {fb['feedback']}", 5)
            pdf.ln(1)

        return bytes(pdf.output())

    except Exception as exc:                            # noqa: BLE001
        logger.error("build_interview_pdf gagal: %s", exc)
        return None
